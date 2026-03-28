from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any
from zipfile import BadZipFile

from docx import Document

from backend.core.config import PROJECT_ROOT


class StageDocumentService:
    def __init__(self) -> None:
        self.base_dir = PROJECT_ROOT / "data" / "stage_docs"
        self.base_dir.mkdir(parents=True, exist_ok=True)

    def build_stage_document(
        self,
        session_id: str,
        project_title: str,
        stage_index: int,
        stage_label: str,
        answers: list[str],
        latest_input: str | None,
        summary: str,
        feedback: str,
        guidance: str,
        draft: str,
        previous_text: str | None = None,
    ) -> dict[str, Any]:
        del feedback, guidance

        file_name = self._file_name(
            session_id=session_id,
            stage_index=stage_index,
            stage_label=stage_label,
            suffix="generated",
        )
        file_path = self.base_dir / file_name

        document = Document()
        if previous_text and previous_text.strip():
            for block in self._split_blocks(previous_text):
                document.add_paragraph(block)
        else:
            document.add_heading(project_title, level=1)
        document.add_paragraph(f"Stage {stage_index} - {stage_label}")

        body = self._build_document_body(
            draft=draft,
            summary=summary,
            answers=answers,
            latest_input=latest_input,
        )
        for block in self._split_blocks(body):
            document.add_paragraph(block)

        document.save(file_path)
        preview_text = self.extract_text(file_path)
        return {
            "stage_index": stage_index,
            "stage_label": stage_label,
            "file_name": file_name,
            "generated_path": str(file_path),
            "latest_path": str(file_path),
            "generated_text": preview_text,
            "latest_text": preview_text,
            "preview_text": preview_text,
            "is_modified": False,
            "modification_summary": None,
            "source": "generated",
            "updated_at": self._timestamp(),
        }

    def save_uploaded_revision(
        self,
        session_id: str,
        stage_index: int,
        stage_label: str,
        file_bytes: bytes,
    ) -> dict[str, Any]:
        file_name = self._file_name(
            session_id=session_id,
            stage_index=stage_index,
            stage_label=stage_label,
            suffix="revised",
        )
        file_path = self.base_dir / file_name
        file_path.write_bytes(file_bytes)

        preview_text = self.extract_text(file_path)
        return {
            "latest_path": str(file_path),
            "latest_text": preview_text,
            "preview_text": preview_text,
            "source": "uploaded",
            "updated_at": self._timestamp(),
        }

    def save_text_revision(
        self,
        session_id: str,
        stage_index: int,
        stage_label: str,
        content: str,
    ) -> dict[str, Any]:
        file_name = self._file_name(
            session_id=session_id,
            stage_index=stage_index,
            stage_label=stage_label,
            suffix="edited",
        )
        file_path = self.base_dir / file_name

        document = Document()
        for block in self._split_blocks(content):
            document.add_paragraph(block)

        document.save(file_path)
        preview_text = self.extract_text(file_path)
        return {
            "latest_path": str(file_path),
            "latest_text": preview_text,
            "preview_text": preview_text,
            "source": "edited",
            "updated_at": self._timestamp(),
        }

    def extract_text(self, file_path: str | Path) -> str:
        path = Path(file_path)
        try:
            document = Document(path)
        except (BadZipFile, KeyError, ValueError) as exc:
            raise ValueError("Could not read that docx file. Please upload a standard Word document.") from exc

        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        return "\n\n".join(paragraphs).strip()

    def download_path(self, metadata: dict[str, Any]) -> Path:
        return Path(str(metadata["latest_path"]))

    def _file_name(
        self,
        session_id: str,
        stage_index: int,
        stage_label: str,
        suffix: str,
    ) -> str:
        safe_label = "".join(char if char.isalnum() else "_" for char in stage_label).strip("_")
        safe_label = safe_label or f"stage_{stage_index}"
        return f"{session_id}_stage_{stage_index:02d}_{safe_label}_{suffix}.docx"

    def _timestamp(self) -> str:
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    def _build_document_body(
        self,
        draft: str,
        summary: str,
        answers: list[str],
        latest_input: str | None,
    ) -> str:
        cleaned_draft = self._clean_document_text(draft)
        if cleaned_draft:
            return cleaned_draft

        answer_fallback = self._compose_answer_fallback(answers, latest_input)
        if answer_fallback:
            return answer_fallback

        cleaned_summary = self._clean_document_text(summary)
        if cleaned_summary:
            return cleaned_summary

        return "This stage still needs more detail."

    def _compose_answer_fallback(self, answers: list[str], latest_input: str | None) -> str:
        blocks = [answer.strip() for answer in answers if answer and answer.strip()]
        if latest_input and latest_input.strip():
            blocks.append(latest_input.strip())
        return "\n".join(blocks)

    def _clean_document_text(self, text: str) -> str:
        if not text or not text.strip():
            return ""

        cleaned = text.replace("\r\n", "\n").strip()

        for pattern in (
            r"^(?:\u8fd9(?:\u4e00)?\u6bb5|\u672c\u9636\u6bb5|\u8fd9\u4e00\u90e8\u5206)(?:\u5148)?(?:\u53ef\u4ee5|\u53ef\u5148)?\u8fd9\u6837\u5199[:\uff1a]\s*",
            r"^(?:\u9636\u6bb5\u8349\u7a3f|\u8349\u7a3f|\u6587\u7a3f|\u9636\u6bb5\u6587\u7a3f)[:\uff1a]\s*",
            r"^(?:draft|working draft|stage draft|document draft)[:\uff1a]\s*",
            r"^(?:you could phrase it this way|you can write it like this|for this stage)[:\uff1a]\s*",
        ):
            cleaned = re.sub(pattern, "", cleaned, count=1)

        for marker in (
            "\u76ee\u524d\u5148\u6536\u4f4f\u8fd9\u51e0\u4e2a\u70b9\uff1a",
            "\u76ee\u524d\u5148\u6536\u4f4f\u8fd9\u51e0\u4e2a\u70b9:",
            "\u5148\u53ef\u4ee5\u8fd9\u6837\u5199\uff1a",
            "\u5148\u53ef\u4ee5\u8fd9\u6837\u5199:",
            "\u53ef\u4ee5\u8fd9\u6837\u5199\uff1a",
            "\u53ef\u4ee5\u8fd9\u6837\u5199:",
            "Here is a working draft:",
            "Here is the working draft:",
            "Draft:",
            "Working draft:",
        ):
            if marker in cleaned:
                _, suffix = cleaned.split(marker, 1)
                if suffix.strip():
                    cleaned = suffix.strip()
                    break

        cleaned = self._strip_keyword_prefix(cleaned)

        ignored_prefixes = (
            "\u7528\u6237\u56de\u7b54",
            "\u8865\u5145\u8bf4\u660e",
            "ai\u9636\u6bb5\u603b\u7ed3",
            "ai\u53cd\u9988",
            "\u9636\u6bb5\u53cd\u9988",
            "\u4e0b\u4e00\u6b65\u5efa\u8bae",
            "\u6307\u5bfc\u5efa\u8bae",
            "coach_message",
            "current_focus_summary",
            "guidance",
            "user answer",
            "user answers",
            "additional note",
            "stage feedback",
            "next step",
            "next steps",
            "working draft",
            "draft",
        )
        lines: list[str] = []
        for raw_line in cleaned.splitlines():
            line = raw_line.strip()
            if not line:
                continue

            plain_line = line.lstrip("-* ").strip()
            if any(plain_line.lower().startswith(prefix) for prefix in ignored_prefixes):
                continue
            lines.append(line)

        return "\n".join(lines).strip()

    def _strip_keyword_prefix(self, text: str) -> str:
        for separator in ("\uff1a", ":"):
            if separator not in text:
                continue

            prefix, suffix = text.split(separator, 1)
            prefix = prefix.strip()
            suffix = suffix.strip()
            if not suffix:
                continue

            if any(
                keyword in prefix
                for keyword in (
                    "\u8fd9\u6837\u5199",
                    "\u8fd9\u4e00\u6bb5",
                    "\u8fd9\u6bb5",
                    "\u8349\u7a3f",
                    "\u6587\u7a3f",
                    "draft",
                    "working draft",
                    "stage draft",
                )
            ):
                return suffix

        return text

    def _split_blocks(self, content: str) -> list[str]:
        blocks = [block.strip() for block in re.split(r"\n\s*\n", content) if block.strip()]
        if blocks:
            return blocks
        return [content.strip()] if content.strip() else []
