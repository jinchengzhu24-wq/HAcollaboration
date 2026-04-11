from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any
from zipfile import BadZipFile

from docx import Document

from backend.config import PROJECT_ROOT


class StageDocumentService:
    def __init__(self) -> None:
        self.base_dir = PROJECT_ROOT / "data" / "stage_docs"
        self.base_dir.mkdir(parents=True, exist_ok=True)

    def build_stage_document(
        self,
        session_id: str,
        project_title: str,
        stage_index: int,
        stage_label: str,
        questions: list[str],
        answers: list[str],
        latest_input: str | None,
        summary: str,
        feedback: str,
        guidance: str,
        draft: str,
    ) -> dict[str, Any]:
        file_name = self._file_name(
            session_id=session_id,
            stage_index=stage_index,
            stage_label=stage_label,
            suffix="generated",
        )
        file_path = self.base_dir / file_name

        document = Document()
        document.add_heading(project_title, level=1)
        document.add_heading(f"Stage {stage_index}: {stage_label}", level=2)

        body = self._build_document_body(
            draft=draft,
            summary=summary,
            feedback=feedback,
            guidance=guidance,
            questions=questions,
            answers=answers,
            latest_input=latest_input,
        )
        for block in self._split_blocks(body):
            document.add_paragraph(block)

        document.save(file_path)
        preview_text = self.extract_text(file_path)
        return {
            "stage_index": stage_index,
            "stage_label": stage_label,
            "file_name": file_name,
            "generated_path": str(file_path),
            "latest_path": str(file_path),
            "generated_text": preview_text,
            "latest_text": preview_text,
            "preview_text": preview_text,
            "is_modified": False,
            "modification_summary": None,
            "source": "generated",
            "updated_at": self._timestamp(),
        }

    def save_uploaded_revision(
        self,
        session_id: str,
        stage_index: int,
        stage_label: str,
        file_bytes: bytes,
    ) -> dict[str, Any]:
        file_name = self._file_name(
            session_id=session_id,
            stage_index=stage_index,
            stage_label=stage_label,
            suffix="revised",
        )
        file_path = self.base_dir / file_name
        file_path.write_bytes(file_bytes)

        preview_text = self.extract_text(file_path)
        return {
            "file_name": file_name,
            "latest_path": str(file_path),
            "latest_text": preview_text,
            "preview_text": preview_text,
            "source": "uploaded",
            "updated_at": self._timestamp(),
        }

    def save_text_revision(
        self,
        session_id: str,
        stage_index: int,
        stage_label: str,
        content: str,
    ) -> dict[str, Any]:
        file_name = self._file_name(
            session_id=session_id,
            stage_index=stage_index,
            stage_label=stage_label,
            suffix="edited",
        )
        file_path = self.base_dir / file_name

        document = Document()
        for block in self._split_blocks(content):
            document.add_paragraph(block)

        document.save(file_path)
        preview_text = self.extract_text(file_path)
        return {
            "file_name": file_name,
            "latest_path": str(file_path),
            "latest_text": preview_text,
            "preview_text": preview_text,
            "source": "edited",
            "updated_at": self._timestamp(),
        }

    def extract_text(self, file_path: str | Path) -> str:
        path = Path(file_path)
        try:
            document = Document(path)
        except (BadZipFile, KeyError, ValueError) as exc:
            raise ValueError("Could not read that docx file. Please upload a standard Word document.") from exc

        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        return "\n\n".join(paragraphs).strip()

    def download_path(self, metadata: dict[str, Any]) -> Path:
        return Path(str(metadata["latest_path"]))

    def save_combined_document_to_desktop(
        self,
        project_title: str,
        sections: list[dict[str, Any]],
    ) -> Path:
        desktop = Path.home() / "Desktop"
        desktop.mkdir(parents=True, exist_ok=True)
        file_path = desktop / f"{self._safe_file_stem(project_title)}_CAR_document.docx"

        document = Document()
        document.add_heading(project_title, level=1)

        for section in sections:
            document.add_heading(f"Stage {section['stage_index']}: {section['stage_label']}", level=2)
            document.add_paragraph("Working Draft:")
            document.add_paragraph(section.get("draft", "") or "This stage still needs more detail.")

            feedback = str(section.get("feedback", "") or "").strip()
            if feedback:
                document.add_paragraph("System Feedback:")
                document.add_paragraph(feedback)

            guidance = str(section.get("guidance", "") or "").strip()
            if guidance:
                document.add_paragraph("Next Step:")
                document.add_paragraph(guidance)

            questions = [str(question).strip() for question in section.get("questions", []) if str(question).strip()]
            if questions:
                document.add_paragraph("Current Questions:")
                for index, question in enumerate(questions, start=1):
                    document.add_paragraph(f"{index}. {question}")

        document.save(file_path)
        return file_path

    def _file_name(
        self,
        session_id: str,
        stage_index: int,
        stage_label: str,
        suffix: str,
    ) -> str:
        safe_label = "".join(char if char.isalnum() else "_" for char in stage_label).strip("_")
        safe_label = safe_label or f"stage_{stage_index}"
        return f"{session_id}_stage_{stage_index:02d}_{safe_label}_{suffix}.docx"

    def _safe_file_stem(self, title: str) -> str:
        safe = "".join(char if char.isalnum() else "_" for char in title).strip("_")
        return safe or "action_research"

    def _timestamp(self) -> str:
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    def _build_document_body(
        self,
        draft: str,
        summary: str,
        feedback: str,
        guidance: str,
        questions: list[str],
        answers: list[str],
        latest_input: str | None,
    ) -> str:
        cleaned_draft = self._clean_document_text(draft)
        answer_fallback = self._compose_answer_fallback(answers, latest_input)
        cleaned_summary = self._clean_document_text(summary)
        working_draft = cleaned_draft or answer_fallback or cleaned_summary or "This stage still needs more detail."

        blocks = ["Working Draft:", working_draft]

        cleaned_feedback = self._clean_document_text(feedback)
        if cleaned_feedback:
            blocks.extend(["System Feedback:", cleaned_feedback])

        cleaned_guidance = self._clean_document_text(guidance)
        if cleaned_guidance:
            blocks.extend(["Next Step:", cleaned_guidance])

        cleaned_questions = [question.strip() for question in questions if question and question.strip()]
        if cleaned_questions:
            blocks.append("Current Questions:")
            blocks.extend(f"{index}. {question}" for index, question in enumerate(cleaned_questions, start=1))

        return "\n\n".join(block for block in blocks if block.strip())

    def _compose_answer_fallback(self, answers: list[str], latest_input: str | None) -> str:
        blocks = [answer.strip() for answer in answers if answer and answer.strip()]
        if latest_input and latest_input.strip():
            blocks.append(latest_input.strip())
        return "\n".join(blocks)

    def _clean_document_text(self, text: str) -> str:
        if not text or not text.strip():
            return ""

        cleaned = text.replace("\r\n", "\n").strip()
        for pattern in (
            r"^(?:这一段|本阶段|这一部分)(?:先)?(?:可以|可先)?这样写[:：]\s*",
            r"^(?:阶段草稿|草稿|文稿|阶段文稿)[:：]\s*",
            r"^(?:draft|working draft|stage draft|document draft)[:：]\s*",
        ):
            cleaned = re.sub(pattern, "", cleaned, count=1, flags=re.IGNORECASE)

        for marker in (
            "目前先收住这几个点：",
            "先可以这样写：",
            "可以这样写：",
            "Here is a working draft:",
            "Here is the working draft:",
            "Draft:",
            "Working draft:",
        ):
            if marker in cleaned:
                _, suffix = cleaned.split(marker, 1)
                if suffix.strip():
                    cleaned = suffix.strip()
                    break

        return cleaned.strip()

    def _split_blocks(self, content: str) -> list[str]:
        blocks = [block.strip() for block in re.split(r"\n\s*\n", content) if block.strip()]
        if blocks:
            return blocks
        return [content.strip()] if content.strip() else []
