from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from backend.main import app
from backend.routes.dialogue import dialogue_service
from backend.services.dialogue_service import DialogueService


@pytest.fixture(autouse=True)
def disable_deepseek_for_route_tests():
    original_client = dialogue_service.deepseek_client
    dialogue_service.deepseek_client = None
    yield
    dialogue_service.deepseek_client = original_client


def test_car_stage_plan_is_returned_on_create() -> None:
    client = TestClient(app)

    response = client.post(
        "/dialogue/sessions",
        json={
            "initial_idea": "I want to improve questioning quality in reading lessons.",
            "project_title": "Questioning Improvement",
        },
    )
    assert response.status_code == 200

    payload = response.json()
    stages = payload["stages"]
    assert len(stages) == 4
    assert payload["active_stage_index"] == 1
    assert payload["current_questions"]
    assert stages[0]["status"] == "available"
    assert stages[0]["is_active"] is True
    assert stages[1]["status"] == "locked"
    assert stages[2]["status"] == "locked"
    assert stages[3]["status"] == "locked"
    assert "Stage 1: Problem Framing" in payload["combined_document"]
    assert "Stage 4: Reflection and Iteration" in payload["combined_document"]


def test_opening_message_responds_to_initial_idea() -> None:
    client = TestClient(app)

    response = client.post(
        "/dialogue/sessions",
        json={
            "initial_idea": "Students give short answers during reading discussions.",
            "project_title": "Reading Discussion",
        },
    )
    assert response.status_code == 200

    opening_message = response.json()["opening_message"]
    assert "Students give short answers during reading discussions." in opening_message
    assert "Starting point captured" not in opening_message

    questions = response.json()["current_questions"]
    assert any("short answers during reading discussions" in question for question in questions)
    assert not any(question.startswith("What is the main classroom problem") for question in questions)


def test_create_session_uses_deepseek_for_initial_reply() -> None:
    class FakeDeepSeekClient:
        def generate_json(self, system_prompt: str, user_prompt: str):
            assert "opening_message" in system_prompt
            assert "Students give one-word answers" in user_prompt
            return {
                "opening_message": (
                    "That sounds like a discussion depth problem, not just a participation problem. "
                    "Let's locate the moment where answers become too short."
                ),
                "questions": [
                    "During which reading task do students most often give one-word answers?",
                    "Which students are most affected, and what longer response would count as early progress?",
                ],
            }

    service = DialogueService()
    service.deepseek_client = FakeDeepSeekClient()

    session = service.create_session(
        project_title="Reading Discussion",
        initial_idea="Students give one-word answers during reading discussions.",
    )

    assert session.state_snapshot["opening_message"].startswith("That sounds like a discussion depth problem")
    assert session.stages[0].questions == [
        "During which reading task do students most often give one-word answers?",
        "Which students are most affected, and what longer response would count as early progress?",
    ]


def test_cida_mode_adds_support_prompts_on_create() -> None:
    client = TestClient(app)

    response = client.post(
        "/dialogue/sessions",
        json={
            "initial_idea": "I want to improve questioning quality in reading lessons.",
            "project_title": "Questioning Improvement",
            "cida_enabled": True,
        },
    )
    assert response.status_code == 200

    payload = response.json()
    first_stage = payload["stages"][0]
    assert payload["cida_enabled"] is True
    assert len(first_stage["cida_guidance"]) == 3
    assert any("Inquiry process" in question for question in payload["current_questions"])
    assert "CIDA support is on" in payload["opening_message"]


def test_cida_mode_can_be_toggled_for_an_existing_session() -> None:
    client = TestClient(app)
    session_id = _create_session(client)

    toggle_response = client.post(
        f"/dialogue/sessions/{session_id}/cida",
        json={"enabled": True},
    )
    assert toggle_response.status_code == 200
    toggle_payload = toggle_response.json()
    assert toggle_payload["cida_enabled"] is True
    assert len(toggle_payload["stages"][0]["cida_guidance"]) == 3
    assert any("questioning quality" in question for question in toggle_payload["current_questions"])
    assert any("Technological support" in question for question in toggle_payload["current_questions"])

    turn_response = client.post(
        f"/dialogue/sessions/{session_id}/stages/1/turn",
        json={
            "answers": [
                "Students often give short answers and the discussion stops too early.",
                "I want to redesign the teacher's follow-up questions.",
            ],
            "latest_input": "The target group is grade seven reading classes.",
        },
    )
    assert turn_response.status_code == 200
    assert "CIDA Support Notes" in turn_response.json()["stages"][0]["draft"]

    off_response = client.post(
        f"/dialogue/sessions/{session_id}/cida",
        json={"enabled": False},
    )
    assert off_response.status_code == 200
    off_payload = off_response.json()
    assert off_payload["cida_enabled"] is False
    assert off_payload["stages"][0]["cida_guidance"] == []
    assert off_payload["stages"][0]["is_outdated"] is True


def test_stage_feedback_is_structured_with_examples() -> None:
    client = TestClient(app)
    session_id = _create_session(client)

    response = client.post(
        f"/dialogue/sessions/{session_id}/stages/1/turn",
        json={
            "answers": [
                "Students often give short answers and the discussion stops too early.",
                "I want to redesign the teacher's follow-up questions.",
            ],
            "latest_input": "The target group is grade seven reading classes.",
        },
    )
    assert response.status_code == 200

    feedback = response.json()["stages"][0]["feedback"]
    assert "What Works:" in feedback
    assert "Make It Sharper:" in feedback
    assert "Example To Add:" in feedback
    assert "Try a sentence like:" in feedback


def test_duplicate_empty_draft_headings_are_collapsed() -> None:
    class FakeDeepSeekClient:
        def generate_json(self, system_prompt: str, user_prompt: str):
            if "opening_message" in system_prompt:
                return {
                    "opening_message": "Let's frame the phone-use issue in one concrete classroom moment.",
                    "questions": [
                        "When is phone use most visible during direct instruction?",
                        "Which students are most affected, and what early change would show progress?",
                    ],
                }
            return {
                "summary": "Problem Framing: phone use during direct instruction.",
                "feedback": (
                    "What Works:\n"
                    "The problem is visible in a specific lesson moment.\n\n"
                    "Make It Sharper:\n"
                    "Name the exact learner group and behavior.\n\n"
                    "Example To Add:\n"
                    "For example, track off-task phone use during the first ten minutes."
                ),
                "guidance": "Clarify **one observable behavior** before confirming this stage.",
                "draft": (
                    "Classroom Problem:\n\n"
                    "Classroom Problem: During direct instruction, many students talk and use their phones.\n\n"
                    "Learner Group:\n\n"
                    "Learner Group: Grade 9 mathematics class of 30 students.\n\n"
                    "Desired Early Change:\n\n"
                    "Desired Early Change: Increase on-task behavior during the first 10 minutes."
                ),
            }

    service = DialogueService()
    service.deepseek_client = FakeDeepSeekClient()
    session = service.create_session(
        project_title="Phone Use",
        initial_idea="Students use phones during direct instruction.",
    )

    service.turn_stage(
        session,
        1,
        ["Many students talk and use their phones during direct instruction."],
        "The target group is a grade 9 mathematics class.",
    )

    draft = session.stages[0].draft or ""
    assert draft.count("Classroom Problem:") == 1
    assert draft.count("Learner Group:") == 1
    assert draft.count("Desired Early Change:") == 1
    assert "Classroom Problem: During direct instruction" in draft


def test_regenerate_after_workspace_edit_updates_feedback_only() -> None:
    client = TestClient(app)
    session_id = _create_session(client)

    turn_response = client.post(
        f"/dialogue/sessions/{session_id}/stages/1/turn",
        json={
            "answers": ["Students use phones during direct instruction."],
            "latest_input": "The target group is grade 10 mathematics.",
        },
    )
    assert turn_response.status_code == 200

    original_draft = turn_response.json()["stages"][0]["draft"]
    edited_draft = (
        f"{original_draft}\n\n"
        "Teacher revision: the clearest problem is off-task phone use in the back row."
    )

    regenerate_response = client.post(
        f"/dialogue/sessions/{session_id}/stages/1/regenerate",
        json={"content": edited_draft},
    )
    assert regenerate_response.status_code == 200

    payload = regenerate_response.json()
    stage = payload["stages"][0]
    assert "I noticed you changed the workspace draft" in payload["message"]
    assert "regenerated the feedback only" in payload["message"]
    assert "Updated Feedback:" in payload["message"]
    assert "What Works:" in payload["message"]
    assert stage["draft"] == edited_draft
    assert "What Works:" in stage["feedback"]


def test_regenerate_without_workspace_edit_updates_draft_and_feedback() -> None:
    class FakeDeepSeekClient:
        def __init__(self) -> None:
            self.stage_payload_count = 0

        def generate_json(self, system_prompt: str, user_prompt: str):
            if "opening_message" in system_prompt:
                return {
                    "opening_message": "Let's frame the phone-use issue.",
                    "questions": [
                        "When is phone use most visible?",
                        "Which students are most affected?",
                    ],
                }
            if "feedback only" in system_prompt:
                return {
                    "feedback": (
                        "What Changed:\n"
                        "The edited draft is clearer.\n\n"
                        "Make It Stronger:\n"
                        "Keep **one observable behavior** central.\n\n"
                        "Example To Add:\n"
                        "For example, tally phone use during direct instruction."
                    )
                }

            self.stage_payload_count += 1
            return {
                "summary": f"Summary version {self.stage_payload_count}",
                "feedback": (
                    "What Works:\n"
                    f"Feedback version **{self.stage_payload_count}** is specific.\n\n"
                    "Make It Sharper:\n"
                    "Name the lesson moment.\n\n"
                    "Example To Add:\n"
                    "For example, record phone checks during instruction."
                ),
                "guidance": "Keep **one observable behavior** central.",
                "draft": (
                    "Classroom Problem:\n"
                    f"Draft version {self.stage_payload_count}.\n\n"
                    "Learner Group:\n"
                    "Grade 10 students.\n\n"
                    "Desired Early Change:\n"
                    "More on-task behavior."
                ),
            }

    service = DialogueService()
    service.deepseek_client = FakeDeepSeekClient()
    session = service.create_session(
        project_title="Phone Use",
        initial_idea="Students use phones during direct instruction.",
    )
    service.turn_stage(
        session,
        1,
        ["Students use phones during direct instruction."],
        "The target group is grade 10 mathematics.",
    )
    first_draft = session.stages[0].draft
    first_feedback = session.stages[0].feedback

    message = service.regenerate_stage(session, 1, first_draft)

    assert "Stage 1 regenerated." in message
    assert "Draft version 2." in (session.stages[0].draft or "")
    assert session.stages[0].draft != first_draft
    assert session.stages[0].feedback != first_feedback


def test_locked_stage_requires_confirm_or_skip_before_unlock() -> None:
    client = TestClient(app)
    session_id = _create_session(client)

    locked_response = client.post(f"/dialogue/sessions/{session_id}/stages/2/activate")
    assert locked_response.status_code == 400

    turn_response = client.post(
        f"/dialogue/sessions/{session_id}/stages/1/turn",
        json={
            "answers": [
                "Students often give short answers and the discussion stops too early.",
                "I want to redesign the teacher's follow-up questions.",
            ],
            "latest_input": "The target group is grade seven reading classes.",
        },
    )
    assert turn_response.status_code == 200
    turn_payload = turn_response.json()
    assert turn_payload["stages"][0]["needs_confirmation"] is True

    confirm_response = client.post(f"/dialogue/sessions/{session_id}/stages/1/confirm")
    assert confirm_response.status_code == 200
    confirm_payload = confirm_response.json()
    assert confirm_payload["active_stage_index"] == 2
    assert confirm_payload["stages"][0]["status"] == "completed"
    assert confirm_payload["stages"][1]["status"] == "available"

    delete_response = client.post(f"/dialogue/sessions/{session_id}/stages/2/delete")
    assert delete_response.status_code == 200
    delete_payload = delete_response.json()
    assert delete_payload["stages"][1]["status"] == "deleted"
    assert delete_payload["stages"][2]["status"] == "available"
    assert delete_payload["active_stage_index"] == 3


def test_deleting_a_future_stage_does_not_unlock_later_stages_too_early() -> None:
    client = TestClient(app)
    session_id = _create_session(client)

    _turn_and_confirm_stage(
        client,
        session_id,
        1,
        [
            "Students often give short answers and the discussion stops too early.",
            "I want to redesign the teacher's follow-up questions.",
        ],
        "The target group is grade seven reading classes.",
    )

    delete_response = client.post(f"/dialogue/sessions/{session_id}/stages/3/delete")
    assert delete_response.status_code == 200
    delete_payload = delete_response.json()
    assert delete_payload["active_stage_index"] == 2
    assert delete_payload["stages"][1]["status"] == "available"
    assert delete_payload["stages"][2]["status"] == "deleted"
    assert delete_payload["stages"][3]["status"] == "locked"


def test_jumping_back_marks_later_stage_outdated_and_confirm_returns_to_current_stage() -> None:
    client = TestClient(app)
    session_id = _create_session(client)

    _turn_and_confirm_stage(
        client,
        session_id,
        1,
        [
            "Students often give short answers and the discussion stops too early.",
            "I want to redesign the teacher's follow-up questions.",
        ],
        "The target group is grade seven reading classes.",
    )
    _turn_and_confirm_stage(
        client,
        session_id,
        2,
        [
            "I plan to model better follow-up prompts during reading discussions.",
            "I will try this first in one grade seven class next week.",
        ],
        "I also want to keep the steps manageable for the first trial.",
    )

    jump_response = client.post(f"/dialogue/sessions/{session_id}/stages/1/activate")
    assert jump_response.status_code == 200
    jump_payload = jump_response.json()
    assert jump_payload["active_stage_index"] == 1
    assert jump_payload["stages"][1]["is_outdated"] is True

    confirm_response = client.post(f"/dialogue/sessions/{session_id}/stages/1/confirm")
    assert confirm_response.status_code == 200
    confirm_payload = confirm_response.json()
    assert confirm_payload["active_stage_index"] == 3


def test_combined_document_edit_marks_later_stage_outdated_until_reviewed() -> None:
    client = TestClient(app)
    session_id = _create_session(client)

    _turn_and_confirm_stage(
        client,
        session_id,
        1,
        [
            "Students often give short answers and the discussion stops too early.",
            "I want to redesign the teacher's follow-up questions.",
        ],
        "The target group is grade seven reading classes.",
    )
    _turn_and_confirm_stage(
        client,
        session_id,
        2,
        [
            "I plan to model better follow-up prompts during reading discussions.",
            "I will try this first in one grade seven class next week.",
        ],
        "I also want to keep the steps manageable for the first trial.",
    )

    session_response = client.get(f"/dialogue/sessions/{session_id}")
    assert session_response.status_code == 200
    combined_document = session_response.json()["combined_document"]
    revised_document = combined_document.replace(
        "Stage 1: Problem Framing\n",
        (
            "Stage 1: Problem Framing\n"
            "The teacher refined the problem statement around follow-up questioning.\n"
            "The revised focus is on extending student talk in grade seven reading lessons."
        ),
        1,
    )

    edit_response = client.post(
        f"/dialogue/sessions/{session_id}/document/edit",
        json={"content": revised_document},
    )
    assert edit_response.status_code == 200
    edit_payload = edit_response.json()
    assert edit_payload["stages"][0]["status"] == "completed"
    assert edit_payload["stages"][1]["is_outdated"] is True
    assert (
        "The teacher refined the problem statement around follow-up questioning."
        in edit_payload["combined_document"]
    )

    review_response = client.post(f"/dialogue/sessions/{session_id}/stages/2/review")
    assert review_response.status_code == 200
    review_payload = review_response.json()
    assert review_payload["stages"][1]["is_outdated"] is False


def test_stage_document_upload_and_download_flow() -> None:
    client = TestClient(app)
    session_id = _create_session(client)

    turn_response = client.post(
        f"/dialogue/sessions/{session_id}/stages/1/turn",
        json={
            "answers": [
                "Students often give short answers and the discussion stops too early.",
                "I want to redesign the teacher's follow-up questions.",
            ],
            "latest_input": "The target group is grade seven reading classes.",
        },
    )
    assert turn_response.status_code == 200
    turn_payload = turn_response.json()
    stage_one = turn_payload["stages"][0]
    download_url = stage_one["document"]["download_url"]
    assert "System Feedback:" in stage_one["document"]["preview_text"]
    assert "Next Step:" in stage_one["document"]["preview_text"]
    assert "Current Questions:" in stage_one["document"]["preview_text"]

    download_response = client.get(download_url)
    assert download_response.status_code == 200
    assert download_response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    revised_doc = Document()
    revised_doc.add_heading("Questioning Improvement", level=1)
    revised_doc.add_paragraph("Stage 1: Problem Framing")
    revised_doc.add_paragraph("Added a revision about follow-up questioning in reading lessons.")
    buffer = BytesIO()
    revised_doc.save(buffer)
    buffer.seek(0)

    upload_response = client.post(
        f"/dialogue/sessions/{session_id}/stages/1/document/upload",
        files={
            "file": (
                "stage_1_revised.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload_response.status_code == 200
    upload_payload = upload_response.json()
    document = upload_payload["stages"][0]["document"]
    assert document["is_modified"] is True
    assert "follow-up questioning" in document["preview_text"]


def test_save_document_exports_a_copy_to_desktop() -> None:
    client = TestClient(app)
    session_id = _create_session(client)

    turn_response = client.post(
        f"/dialogue/sessions/{session_id}/stages/1/turn",
        json={
            "answers": [
                "Students often give short answers and the discussion stops too early.",
                "I want to redesign the teacher's follow-up questions.",
            ],
            "latest_input": "The target group is grade seven reading classes.",
        },
    )
    assert turn_response.status_code == 200
    combined_document = client.get(f"/dialogue/sessions/{session_id}").json()["combined_document"]

    save_response = client.post(
        f"/dialogue/sessions/{session_id}/document/save",
        json={"content": combined_document},
    )
    assert save_response.status_code == 200
    message = save_response.json()["message"]
    assert "Saved a copy to" in message

    save_path = Path(message.split("Saved a copy to ", 1)[1].rstrip("."))
    assert save_path.exists()


def _create_session(client: TestClient) -> str:
    response = client.post(
        "/dialogue/sessions",
        json={
            "initial_idea": "I want to improve questioning quality in reading lessons.",
            "project_title": "Questioning Improvement",
        },
    )
    assert response.status_code == 200
    return response.json()["session_id"]


def _turn_and_confirm_stage(
    client: TestClient,
    session_id: str,
    stage_index: int,
    answers: list[str],
    latest_input: str | None,
) -> None:
    turn_response = client.post(
        f"/dialogue/sessions/{session_id}/stages/{stage_index}/turn",
        json={
            "answers": answers,
            "latest_input": latest_input,
        },
    )
    assert turn_response.status_code == 200

    confirm_response = client.post(f"/dialogue/sessions/{session_id}/stages/{stage_index}/confirm")
    assert confirm_response.status_code == 200
